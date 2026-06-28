from __future__ import annotations

from io import BytesIO
from typing import BinaryIO

from docx import Document
from pypdf import PdfReader


def read_docx(file_bytes: bytes) -> str:
    """讀取 Word 教案文字。"""
    doc = Document(BytesIO(file_bytes))
    chunks: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                text = cell.text.strip().replace("\n", " ")
                if text and (not cells or cells[-1] != text):
                    cells.append(text)
            if any(cells):
                chunks.append(" | ".join(cells))

    return "\n".join(chunks).strip()


def read_pdf(file_bytes: bytes) -> str:
    """讀取可選取文字的 PDF。掃描圖檔 PDF 需要先轉成可辨識文字。"""
    reader = PdfReader(BytesIO(file_bytes))
    chunks: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        text = text.strip()
        if text:
            chunks.append(f"[第 {i} 頁]\n{text}")
    return "\n\n".join(chunks).strip()


def read_txt(file_bytes: bytes) -> str:
    """讀取純文字檔，優先使用 UTF-8，失敗再嘗試常見中文編碼。"""
    for encoding in ("utf-8-sig", "utf-8", "cp950", "big5"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore").strip()


def read_uploaded_file(uploaded_file: BinaryIO, file_name: str) -> str:
    data = uploaded_file.read()
    lower_name = file_name.lower()

    if lower_name.endswith(".docx"):
        return read_docx(data)
    if lower_name.endswith(".pdf"):
        return read_pdf(data)
    if lower_name.endswith(".txt"):
        return read_txt(data)

    raise ValueError("目前僅支援 DOCX、PDF、TXT 檔案。")
