from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from schema import CHECKLIST_ITEMS, RATING_OPTIONS, SELF_CHECK_ITEMS


def _set_font(run, size=11, bold=False):
    run.font.name = "標楷體"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    run.font.size = Pt(size)
    run.bold = bold


def _p(cell, text: str = "", bold=False, size=11, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text or ""))
    _set_font(run, size=size, bold=bold)
    return paragraph


def _add_paragraph(doc: Document, text: str = "", bold=False, size=11, align=None):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text or ""))
    _set_font(run, size=size, bold=bold)
    return paragraph


def _cell_shading(cell, fill="D9D9D9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _cell_text(cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs)


def _merge_row_cells(row, start: int, end: int):
    merged = row.cells[start]
    for idx in range(start + 1, end + 1):
        merged = merged.merge(row.cells[idx])
    return merged


def _set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def _setup_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    styles = doc.styles
    styles["Normal"].font.name = "標楷體"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    styles["Normal"].font.size = Pt(11)


def _title(doc: Document, appendix: str, title: str):
    _add_paragraph(doc, appendix, bold=True, size=12)
    p = _add_paragraph(doc, "基隆市 114 學年度學校辦理校長及教師公開授課", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.paragraph_format.space_after = Pt(0)
    _add_paragraph(doc, title, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)


def _basic_table(doc: Document, data: dict[str, Any], include_time=True):
    basic = data["basic_info"]
    table = doc.add_table(rows=4 if include_time else 3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    rows = table.rows
    entries = [
        ("教學時間", basic.get("lesson_minutes", ""), "教學班級", basic.get("class_name", "")),
        ("教學領域", basic.get("subject_area", ""), "教學單元", basic.get("unit_name", "")),
        ("教 學 者", basic.get("teacher", ""), "觀 察 者", basic.get("observer", "")),
    ]
    if include_time:
        entries.append(("觀察時間", basic.get("period_time", ""), "觀察後會談時間", basic.get("post_meeting_minutes", "")))
    for r, values in zip(rows, entries):
        for i, value in enumerate(values):
            _p(r.cells[i], value, bold=i in (0, 2), align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2) else None)
            if i in (0, 2):
                _cell_shading(r.cells[i])
    return table


def add_appendix2(doc: Document, data: dict[str, Any]):
    _title(doc, "附表 2：", "共同備課紀錄表")
    _basic_table(doc, data, include_time=True)
    doc.add_paragraph()
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    fields = ["教材內容", "教學目標", "學生經驗", "教學活動", "教學評量方式", "觀察的工具和觀察焦點"]
    for row, field in zip(table.rows, fields):
        _p(row.cells[0], f"{field}：", bold=True)
        _cell_shading(row.cells[0])
        _p(row.cells[1], data["appendix2"].get(field, ""))
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _add_paragraph(doc, "(本表為參考格式，學校得視需求修改)", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(doc, "授課教師：_____________________    觀課教師：_____________________", size=11)


def _rating_mark(selected: str, option: str) -> str:
    return "■" if selected == option else "□"


def add_appendix3(doc: Document, data: dict[str, Any]):
    doc.add_page_break()
    _title(doc, "附表 3：", "觀課紀錄表")
    _basic_table(doc, data, include_time=True)
    doc.add_paragraph()
    table = doc.add_table(rows=1 + len(CHECKLIST_ITEMS), cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    headers = ["層面", "檢核項目", "檢核重點", "優良", "普通", "可改進", "未呈現"]
    for cell, header in zip(table.rows[0].cells, headers):
        _p(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_shading(cell)

    item_group = {
        "1": "1.清楚呈現教材內容",
        "2": "2.運用有效教學技巧",
        "3": "3.應用良好溝通技巧",
        "4": "4.運用學習評量評估學習成效",
        "5": "5.維持良好的班級秩序以促進學習",
        "6": "6.營造積極的班級氣氛",
    }
    ratings = data["appendix3"].get("ratings", {})
    current_layer = ""
    current_group = ""
    for i, (code, desc) in enumerate(CHECKLIST_ITEMS, start=1):
        row = table.rows[i]
        layer = "教師教學" if code[0] in "1234" else "班級經營"
        group = item_group.get(code[0], "")
        _p(row.cells[0], layer if layer != current_layer else "", align=WD_ALIGN_PARAGRAPH.CENTER)
        _p(row.cells[1], group if group != current_group else "", bold=bool(group != current_group))
        _p(row.cells[2], f"{code} {desc}")
        selected = ratings.get(code, "普通")
        for j, option in enumerate(RATING_OPTIONS, start=3):
            _p(row.cells[j], _rating_mark(selected, option), align=WD_ALIGN_PARAGRAPH.CENTER)
        current_layer = layer
        current_group = group
    _add_paragraph(doc, "(本表為參考格式，學校得視需求修改)", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(doc, "授課教師：_____________________    觀課教師：_____________________", size=11)


def add_appendix4(doc: Document, data: dict[str, Any]):
    doc.add_page_break()
    _title(doc, "附表 4：", "教學自我省思檢核表")
    basic = data["basic_info"]
    _add_paragraph(doc, f"授課教師姓名：{basic.get('teacher','')}    教學班級：{basic.get('class_name','')}    教學領域：{basic.get('subject_area','')}")
    _add_paragraph(doc, f"教學單元名稱：{basic.get('unit_name','')}")
    table = doc.add_table(rows=1 + len(SELF_CHECK_ITEMS), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    headers = ["序號", "檢核項目", "優良", "普通", "可改進", "未呈現"]
    for cell, header in zip(table.rows[0].cells, headers):
        _p(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_shading(cell)
    self_ratings = data["appendix4"].get("self_ratings", {})
    for idx, item in enumerate(SELF_CHECK_ITEMS, start=1):
        row = table.rows[idx]
        _p(row.cells[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
        _p(row.cells[1], item)
        selected = self_ratings.get(item, "普通")
        for j, option in enumerate(RATING_OPTIONS, start=2):
            _p(row.cells[j], _rating_mark(selected, option), align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "◎教學省思：", bold=True)
    _add_paragraph(doc, data["appendix4"].get("教學省思", ""))
    _add_paragraph(doc, "(本表為參考格式，學校得視需求修改)", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(doc, "授課教師：_____________________    觀課教師：_____________________", size=11)


def add_appendix5(doc: Document, data: dict[str, Any]):
    doc.add_page_break()
    _title(doc, "附表 5：", "議課紀錄表")
    _basic_table(doc, data, include_time=True)
    fields = ["教學者教學優點與特色", "教學者教學待調整或改變之處", "對教學者之具體成長建議"]
    for i, field in enumerate(fields, start=1):
        _add_paragraph(doc, f"{['一','二','三'][i-1]}、{field}：", bold=True)
        _add_paragraph(doc, data["appendix5"].get(field, ""))
    _add_paragraph(doc, "(本表為參考格式，學校得視需求修改)", size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(doc, "授課教師：_____________________    觀課教師：_____________________", size=11)


def add_appendix8(doc: Document, data: dict[str, Any]):
    doc.add_page_break()
    _title(doc, "附表 8：", "觀課紀錄表（結構式）")
    basic = data["basic_info"]
    _add_paragraph(doc, f"學校：{basic.get('school','')}    觀課科目：{basic.get('subject_area','')}")
    _add_paragraph(doc, f"授課教師：{basic.get('teacher','')}    觀課班級：{basic.get('class_name','')}")
    _add_paragraph(doc, f"授課單元名稱：{basic.get('unit_name','')}    觀課日期：{basic.get('period_time','')}")
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    fields = ["學生上課狀況", "學生分組討論情形", "知識學習的情形", "綜合建議"]
    prompts = [
        "(1)學生投入課堂學習程度如何？\n(2)學生有干擾課堂行為嗎？情形如何？",
        "(1)小組間互動情形如何？\n(2)小組討論是否聚焦本次課堂？\n(3)小組討論內容深度？",
        "(1)學生在課堂中哪一個部分感到興趣？\n(2)學生在學習中有沒有困難之處？\n(3)真正有效的學習發生在什麼情境？",
        "",
    ]
    for row, field, prompt in zip(table.rows, fields, prompts):
        _p(row.cells[0], f"{field}\n{prompt}", bold=True)
        _cell_shading(row.cells[0])
        _p(row.cells[1], data["appendix8"].get(field, ""))
    _add_paragraph(doc, "觀課人員：______________________________________", size=11)


def add_appendix10(doc: Document, data: dict[str, Any]):
    doc.add_page_break()
    _title(doc, "附表 10：", "學習共同體公開觀課紀錄表")
    basic = data["basic_info"]
    _add_paragraph(doc, f"觀課科目：{basic.get('subject_area','')}    授課教師：{basic.get('teacher','')}    觀課班級：{basic.get('class_name','')}")
    _add_paragraph(doc, f"授課內容：{basic.get('unit_name','')}    觀課日期：{basic.get('period_time','')}    觀課者：{basic.get('observer','')}")
    _add_paragraph(doc, "觀課重點", bold=True)
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    headers = ["1.全班學習氛圍", "2.學生學習動機與歷程", "3.學生學習結果"]
    points = [
        "1-1 是否有安心學習的環境？\n1-2 是否有熱衷學習的環境？\n1-3 是否有聆聽學習的環境？",
        "2-1 老師是否關照每個學生的學習？\n2-2 是否引發學生學習動機？\n2-3 學生學習動機是否持續？\n2-4 學生是否相互關注與傾聽？\n2-5 學生是否互相協助與討論？\n2-6 學生是否投入和參與學習？\n2-7 是否發現特殊表現的學生？",
        "3-1 學生學習如何發生？何時發生？\n3-2 學生學習的困難之處是什麼？\n3-3 學習挑戰是否發生？\n3-4 學生學習思考程度是否深化？\n3-5 學生是否能進行反思或後設思考？",
    ]
    for i in range(3):
        _p(table.rows[0].cells[i], headers[i], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_shading(table.rows[0].cells[i])
        _p(table.rows[1].cells[i], points[i])

    _add_paragraph(doc, "課堂軼事紀錄", bold=True)
    rows = data["appendix10"].get("課堂軼事紀錄", []) or []
    t2 = doc.add_table(rows=1 + len(rows), cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(t2)
    for cell, header in zip(t2.rows[0].cells, ["時間", "教師學習引導", "學生學習行為", "備註"]):
        _p(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_shading(cell)
    for r, rowdata in enumerate(rows, start=1):
        row = t2.rows[r]
        for c, key in enumerate(["時間", "教師學習引導", "學生學習行為", "備註"]):
            _p(row.cells[c], rowdata.get(key, ""))
    _add_paragraph(doc, "觀課的學習", bold=True)
    _add_paragraph(doc, data["appendix10"].get("觀課的學習", ""))


def build_docx(data: dict[str, Any]) -> bytes:
    doc = Document()
    _setup_doc(doc)
    add_appendix2(doc, data)
    add_appendix3(doc, data)
    add_appendix4(doc, data)
    add_appendix5(doc, data)
    output = BytesIO()
    doc.save(output)
    return output.getvalue()
